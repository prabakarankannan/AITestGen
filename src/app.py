from flask import Flask, request, jsonify, render_template, send_file
from docx import Document
from docx.shared import Inches
from flask import send_from_directory
import os
import io
import pandas as pd
from bs4 import BeautifulSoup
from dotenv import load_dotenv
load_dotenv()

# Load the Azure OpenAI API key from the environment variables
from groq import Groq

client = Groq(
    api_key="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",
)

# https://escola42.openai.azure.com/openai/deployments/london_is_best/completions?api-version=2024-02-15-preview

CHAT_COMPLETIONS_MODEL = 'llama3-70b-8192'
SEED=123
OUTPUT_FOLDER = './out'
PROMPT_FOLDER = './prompts'
SYSTEM_MESSAGE = 'system_message_html.txt'
SYSTEM_MESSAGE_AMEND = 'system_message_amend.txt'
STATIC_FOLDER = './static'
TEST_CASE_RESULT_PAGE = 'result.html'

# Load the system message from the file
with open(PROMPT_FOLDER + '/' + SYSTEM_MESSAGE, 'r') as file:
	system_message = file.read()

messages = [""" {"role":"system","content": system_message} """]

app = Flask(__name__)

def extract_text_from_file(file):
    filename = file.filename.lower()
    if filename.endswith('.docx'):
        # Extract text from .docx file
        doc = Document(file)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)
    else:
        # Read text from .txt file
        return file.read().decode('utf-8')

def get_modified_system_message(system_message, automatic_manual, test_type):

	print(system_message)
	print(automatic_manual)
	print(test_type)
	
	# Replace the placeholders in the system message with the actual values
	if automatic_manual == 'automatic':
		system_message = system_message.replace('<AUTOMATIC_MANUAL_TEST>', 'automated')
	elif automatic_manual == 'manual':
		system_message = system_message.replace('<AUTOMATIC_MANUAL_TEST>', 'manual')
	else:
		system_message = system_message.replace('The tests must be generated to be <AUTOMATIC_MANUAL_TEST>.', '')

	if test_type == 'unit':
		system_message = system_message.replace('<TEST_TYPE>', 'unit tests')
	elif test_type == 'integration':
		system_message = system_message.replace('<TEST_TYPE>', 'integration tests')
	elif test_type == 'system':
		system_message = system_message.replace('<TEST_TYPE>', 'system tests')
	else:
		system_message = system_message.replace('The tests must follow the design of <TEST_TYPE>.', '')
	
	print(system_message)
	return system_message

"""
Endpoint to generate test cases from user story
Method: POST
Parameters:	user_story: string or file
Returns:	test_cases: string
"""
@app.route('/generate_tests', methods=['POST'])
def generate_test_cases():

	user_story = ''
	print('yooo')

	if 'file' in request.files:
		file = request.files['file']
		user_story = extract_text_from_file(file)
		if user_story == 'Unsupported file format':
			return jsonify({"error": "Unsupported file format. Please upload a .txt or .docx file"})
		print('got file')
		print(user_story)
	else:
		user_story = request.form.get('text')
		print('got text')
		print(user_story)

	automatic_manual_flag = request.form.get('automaticManual')
	test_type = request.form.get('testType')
	
	if user_story is None or user_story == '':
		return jsonify({"error": "User story not provided"})
	if automatic_manual_flag is None or automatic_manual_flag == '':
		return jsonify({"error": "Automatic/Manual flag not provided"})
	if test_type is None or test_type == '':
		return jsonify({"error": "Test type not provided"})
	
	global messages
	messages = []

	messages.append({"role":"system", "content": get_modified_system_message(system_message, automatic_manual_flag, test_type)})

	#Remove every &nbsp; from the user story
	user_story = user_story.replace('&nbsp;', ' ')
	messages.append({"role":"user","content": user_story})

	# Call the Azure OpenAI API to generate test cases
	response = client.chat.completions.create(
	model="llama3-70b-8192",
		messages = messages,
		temperature=0.5,
		max_tokens=4096,
		top_p=0.95,
		frequency_penalty=0,
		presence_penalty=0,
		stop=None
	)

	# Extracting generated test cases from the completion
	test_cases = response.choices[0].message.content
	data = jsonify({"test_cases": test_cases})

	messages.append({"role":"system", "content": test_cases})

	# write test_cases in a file
	with open(OUTPUT_FOLDER + '/test_cases.txt', 'w') as file:
		file.write(test_cases)

	return test_cases


"""
Endpoint to amend test cases using alreadt generated test cases and a user message
Method: POST
Parameters:	message: string
Returns:	test_cases: string
"""
@app.route('/amend_tests', methods=['POST'])
def amend_test_cases():
	request_data = request.get_json()
	message = request_data.get('amended_text')

	print(request_data)
	print(message)

	if message is None or message == '':
		return jsonify({"error": "Message not provided"})
	
	if messages is None or len(messages) == 0:
		return jsonify({"error": "No test cases to amend"})

	system_message = ''

	with open(PROMPT_FOLDER + '/' + SYSTEM_MESSAGE_AMEND, 'r') as file:
		system_message = file.read()

	messages.append({"role":"user","content": system_message + "'" + message + "'"})

	response = client.chat.completions.create(
		model="llama3-70b-8192",
		messages = messages,
		temperature=0.5,
		max_tokens=4096,
		top_p=0.95,
		frequency_penalty=0,
		presence_penalty=0,
		stop=None
	)

	# Extracting generated test cases from the completion
	test_cases = response.choices[0].message.content
	data = jsonify({"test_cases": test_cases})

	messages.append({"role":"system", "content": test_cases})

		# write test_cases in a file
	with open(OUTPUT_FOLDER + '/test_cases.txt', 'w') as file:
		file.write(test_cases)

	return test_cases

@app.route('/download_excel', methods=['POST'])
def download_excel():
    # Get the HTML content from the request
    html_content = request.data.decode('utf-8')

    # Parse the HTML content using BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')

    # Initialize list to collect test case data
    test_cases = []

    # Extract the story name
    story_name = soup.find('h2').text.replace("Test cases for user story ", "").replace(":", "")

    # Loop through each test case
    for test_case_section in soup.find_all('h3'):
        test_case_name = test_case_section.text.split(":")[1].strip()
        
        # Find the next sibling tables
        tables = test_case_section.find_all_next('table', limit=2)
        
        # Extract first table (test case details)
        details_table = tables[0]
        details_data = {}
        for row in details_table.find_all('tr'):
            cols = row.find_all(['th', 'td'])
            if len(cols) == 2:
                key = cols[0].text.strip()
                value = cols[1].text.strip()
                details_data[key] = value

        # Extract second table (actions and expected results)
        actions_table = tables[1]
        actions = []
        expected_results = []
        for row in actions_table.find_all('tr')[1:]:  # Skip header row
            cols = row.find_all('td')
            if len(cols) == 2:
                actions.append(cols[0].text.strip())
                expected_results.append(cols[1].text.strip())

        # Combine all data into a single dictionary
        test_case_data = {
            "Story Name": story_name,
            "Test Case Name": test_case_name,
            "Title": details_data.get('Title', ''),
            "Description": details_data.get('Description', ''),
            "Pre-Conditions": details_data.get('Pre-Conditions', ''),
            "Requirements": details_data.get('Requirements', ''),
            "Actions": "; ".join(actions),
            "Expected Results": "; ".join(expected_results)
        }

        # Add the test case data to the list
        test_cases.append(test_case_data)

    # Create a DataFrame
    df = pd.DataFrame(test_cases)

    # Create a BytesIO object to hold the Excel file in memory
    output = io.BytesIO()
    writer = pd.ExcelWriter(output, engine='xlsxwriter')
    df.to_excel(writer, index=False, sheet_name='TestCases')
    writer.close()
    output.seek(0)

    return send_file(output, download_name='test_cases.xlsx', as_attachment=True)

"""
Endpoint to download the test cases in .txt format
Method: GET
Parameters:	None
Returns:	test_cases.txt
"""
@app.route('/download_docx', methods=['GET'])
def generate_docx():
	# HTML content is retrieved from the file
	html_content = ''
	with open(OUTPUT_FOLDER + '/test_cases.txt', 'r') as file:
		html_content = file.read()

	# Create a new Document
	doc = Document()

	# Parse HTML content
	soup = BeautifulSoup(html_content, 'html.parser')

	# Stylize tables in the Document
	for table in soup.find_all('table'):
		doc.add_table(rows=len(table.find_all('tr')), cols=len(table.find_all(['th', 'td'])))
		for i, row in enumerate(table.find_all('tr')):
			for j, cell in enumerate(row.find_all(['th', 'td'])):
				doc.tables[-1].cell(i, j).text = cell.get_text()
				doc.tables[-1].cell(i, j).paragraphs[0].style = doc.styles['Normal']

	# Save the Document
	docx_file_path = 'output.docx'
	doc.save(docx_file_path)

	# Send the .docx file to the client
	return send_file(
		docx_file_path,
		mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
		as_attachment=True
	)

"""
Endpoint to clear all messages
Method: POST
Parameters:	None
Returns:	None
"""
@app.route('/clear_messages', methods=['POST'])
def clear_messages():
	global messages
	messages = [{"role":"system", "content": system_message}]
	return jsonify({"message": "Messages cleared"})


"""
Endpoint to return index.html
Parameters:	None
Returns:	index.html
"""
@app.route('/', methods=['GET'])
def get_index():
	return send_from_directory('static', 'home.html')

if __name__ == '__main__':
	app.run(debug=True)
